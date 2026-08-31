from docx import Document
from src.document_parser.base_parser import BaseDocumentParser

class DocxParser(BaseDocumentParser):
    def parse(self, file_path: str) -> str:
        doc = Document(file_path)
        text_list = []
        # 读取段落文字
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text_list.append(para_text)
        # 读取表格内文字（教案表格知识点）
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_list.append(" | ".join(row_text))

        return "\n".join(text_list)